# -*- coding: utf-8 -*-
"""
老年人权益保障法规 · 维度自动抽取脚本 v2
========================================
v2 修正：
  1. 按魔数识别真实格式（新疆 .docx 实为 .doc），不再依赖扩展名
  2. 年龄区间取起点（"八十至八十九周岁" → 80）
  3. 高龄津贴排除 >=100 周岁（那是百岁补贴，不是高龄津贴）
  4. 护理假改用「子句窗口 + 独生子女优先」匹配，并排除"非独生子女"
  5. 区分「本法规未规定」/「有陪护时间但未明确天数」/「有明确天数」

用法：
  <venv>/Scripts/python.exe scripts/extract_elderly_law.py
"""

import json
import re
import unicodedata
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "assets"
OUT = ROOT / "scripts" / "raw_extract.json"

CN_DIGITS = "一二三四五六七八九十百两"
NUM = rf"(\d+|[{CN_DIGITS}]+)"
LEAVE_KW = ["护理假", "陪护假", "照料假"]
CARE_TIME_KW = ["陪护时间", "护理时间", "照料时间"]


# ---------------- 文本抽取（按魔数识别） ----------------

def _keep(ch: str) -> bool:
    o = ord(ch)
    if o < 32 or o == 127:
        return False
    if 0x4E00 <= o <= 0x9FFF:
        return True
    if 0x3000 <= o <= 0x303F:
        return True
    if 0xFF00 <= o <= 0xFFEF:
        return True
    return o < 0x2E80


def read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def read_doc(path: Path) -> str:
    """OLE 复合文档：UTF-16LE 解码 + 中文过滤。"""
    raw = path.read_bytes()
    for enc in ("utf-16-le", "utf-16-be"):
        txt = raw.decode(enc, errors="ignore")
        cleaned = "".join(ch for ch in txt if _keep(ch))
        cjk = sum(1 for ch in cleaned if 0x4E00 <= ord(ch) <= 0x9FFF)
        if cjk > 200:
            return cleaned
    return ""


def reflow(text: str) -> str:
    """
    PDF 抽取的文本带硬换行，会把一整句拆到多行
    （如澳门"长者使用文化、康乐及体育设施"和"依法享有收费优惠"分处两行），
    导致关键词跨行无法同时命中。这里把不以句末标点结尾的行与下一行合并。
    """
    buf, out = "", []
    for ln in text.split("\n"):
        s = ln.strip()
        if not s:
            continue
        buf += s
        if s[-1] in "。；：！？":
            out.append(buf)
            buf = ""
    if buf:
        out.append(buf)
    return "\n".join(out)


def read_pdf(path: Path) -> str:
    import pdfplumber
    out = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return reflow("\n".join(out))


def sniff(path: Path) -> str:
    """按魔数判断真实格式。"""
    head = path.read_bytes()[:8]
    if head[:4] == b"%PDF":
        return "pdf"
    if head[:4] == b"PK\x03\x04":
        return "docx"
    if head[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        return "doc"
    return "unknown"


def to_simplified(text: str) -> str:
    """
    港澳台法规是繁体，且术语与内地不同（長者/老人、敬老金、長期照顧…）。
    统一先转简体再复用内地规则，避免为三地各写一套正则。
    内地简体文本转换后基本不变，无副作用。
    """
    try:
        import zhconv
        return zhconv.convert(text, "zh-cn")
    except Exception:
        return text


def read_any(path: Path) -> str:
    kind = sniff(path)
    # 对 PDF：若同目录有 <stem>_ocr.txt 缓存且有内容，优先使用，
    # 避免每次都跑 OCR（香港《安老院条例》是扫描件，已缓存 OCR 结果到 assets）
    if kind == 'pdf':
        for suffix in ('_ocr.txt', '.ocr.txt'):
            cache = path.with_name(path.stem + suffix)
            if cache.is_file():
                txt = cache.read_text(encoding='utf-8', errors='ignore')
                if len(txt) > 200:
                    return to_simplified(unicodedata.normalize('NFKC', txt))
    try:
        if kind == "docx":
            t = read_docx(path)
        elif kind == "doc":
            t = read_doc(path)
        elif kind == "pdf":
            t = read_pdf(path)
        else:
            t = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        print(f"  [warn] {path.name} ({kind}): {e}")
        # docx 解析失败（可能是伪装扩展名）时回退到 OLE 解析
        if kind == "docx":
            try:
                t = read_doc(path)
            except Exception:
                t = ""
        else:
            t = ""
    return to_simplified(unicodedata.normalize("NFKC", t or ""))


# ---------------- 中文数字 ----------------

_CN = {"零": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4,
       "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}


def cn2num(s: str):
    if not s:
        return None
    s = s.strip()
    if s.isdigit():
        return int(s)
    if "百" in s:
        head, tail = s.split("百", 1)
        h = _CN.get(head, 1) if head else 1
        return h * 100 + (cn2num(tail) or 0 if tail else 0)
    if "十" in s:
        head, tail = s.split("十", 1)
        h = _CN.get(head, 1) if head else 1
        return h * 10 + (_CN.get(tail, 0) if tail else 0)
    v = 0
    for ch in s:
        if ch in _CN:
            v = v * 10 + _CN[ch]
        else:
            return None
    return v if s else None


# ---------------- 数值抽取 ----------------

def extract_age(text: str, max_val: int = 200, min_val: int = 40):
    """先取区间起点（八十至八十九 → 80），再取普通年龄。"""
    m = re.search(rf"{NUM}\s*(?:至|-|—)\s*{NUM}\s*周岁", text)
    if m:
        v = cn2num(m.group(1))
        if v is not None and min_val <= v <= max_val:
            return v
    # 英文 age 模式：the age of 60 years / attained the age of 60（港英版条例用）
    m = re.search(rf"(?:the age of|attained the age of)\s*{NUM}\s*years?", text)
    if m:
        v = cn2num(m.group(1))
        if v is not None and min_val <= v <= max_val:
            return v
    for pat in [rf"{NUM}\s*周岁", rf"年满\s*{NUM}", rf"{NUM}\s*岁"]:
        for m in re.finditer(pat, text):
            v = cn2num(m.group(1))
            if v is not None and min_val <= v <= max_val:
                return v
    return None


DAY_PATS = [rf"{NUM}\s*(?:个)?(?:工作)?日", rf"{NUM}\s*(?:个)?(?:工作)?天"]


def extract_days(text: str):
    for pat in DAY_PATS:
        m = re.search(pat, text)
        if m:
            v = cn2num(m.group(1))
            if v is not None and 1 <= v <= 60:
                return v
    return None


def clauses_of(text: str):
    t = re.sub(r"[ \t\r]+", "", text)
    return [c for c in re.split(r"[，,；;。\n、]", t) if len(c.strip()) >= 3]


# ---------------- 护理假专项抽取 ----------------

def extract_care_leave(texts: dict):
    """
    返回 dict:
      days        : 独生子女护理假天数（优先）
      daysGeneral : 不区分是否独生子女的护理假天数
      paid        : 是否带薪 True/False/None
      status      : 'explicit' 有明确天数 | 'timeonly' 仅规定陪护时间 | 'none' 本法规未规定
      evidence    : 证据句
    """
    res = dict(days=None, daysGeneral=None, paid=None,
               status="none", evidence="", nonOnlyChildDays=None)
    cands = []          # (priority, days, clause, is_paid_ctx)
    has_time_only = False

    for t in texts.values():
        cls = clauses_of(t)
        for i, c in enumerate(cls):
            if not any(k in c for k in LEAVE_KW):
                if any(k in c for k in CARE_TIME_KW):
                    has_time_only = True
                continue
            # 前置窗口放宽到 2 个子句：如广西「独生子女父母年满六十周岁的，患病住院期间，
            # 用人单位应当给予其子女每年累计不超过十五天的护理假」跨了 2 个子句
            prev = "".join(cls[max(0, i - 2):i])
            prev1 = cls[i - 1] if i > 0 else ""
            nxt = cls[i + 1] if i + 1 < len(cls) else ""

            only = ("独生子女" in c or "独生子女" in prev)
            non_only = ("非独生子女" in c or "非独生子女" in prev)
            # 天数优先取本子句，其次紧邻的上一子句（prev1），避免跨句误抓
            days = extract_days(c) or extract_days(prev1) or extract_days(nxt)

            if only and not non_only:
                prio = 1
            elif non_only:
                prio = 3
            else:
                prio = 2
            if days is not None:
                cands.append((prio, days, (prev + "，" + c)[:110], only and not non_only))

    if cands:
        cands.sort(key=lambda x: (x[0], -x[1]))
        top_prio = cands[0][0]
        same = [c for c in cands if c[0] == top_prio]
        days = max(c[1] for c in same)
        ev = same[0][2]
        if top_prio == 1:
            res["days"] = days
        elif top_prio == 2:
            res["daysGeneral"] = days
        else:
            res["nonOnlyChildDays"] = days
        res["evidence"] = ev
        res["status"] = "explicit"
        # 若同时抓到独生子女(1)与其他(2/3)，都记录
        for prio, d, _, _ in cands:
            if prio == 1:
                res["days"] = max(res["days"] or 0, d)
            elif prio == 2:
                res["daysGeneral"] = max(res["daysGeneral"] or 0, d)
            elif prio == 3:
                res["nonOnlyChildDays"] = max(res["nonOnlyChildDays"] or 0, d)
    elif has_time_only:
        res["status"] = "timeonly"

    # 带薪判定
    blob = ""
    for t in texts.values():
        for c in clauses_of(t):
            if any(k in c for k in LEAVE_KW):
                blob += c + "；"
    if blob:
        if any(k in blob for k in ["工资照发", "带薪", "工资福利待遇不变", "视为出勤",
                                   "照发工资", "不扣减工资", "工资、奖金", "薪酬照发"]):
            res["paid"] = True
        elif any(k in blob for k in ["不带薪", "不支付工资", "无薪"]):
            res["paid"] = False
    return res


# ---------------- 其他维度 ----------------

# 术语别名：港澳台与内地用词差异很大（長者/老人、敬老金、長期照顧…）
# 文本已统一转简体，这里只需补简体内的别名
ELDERLY_ALIASES = ["老年人", "长者", "老人"]
# 用于"条款未写年龄 → 按当地老年定义推导"的兜底判定
ELDERLY_TERMS = ("老年人", "长者", "老人", "老龄")

DIM_RULES = [
    dict(key="highAgeAllowanceAge", label="高龄津贴起始年龄", mode="age",
         # allow_fallback=False：本项禁用"按老年定义推导"兜底。
         # 原因：锚点含"长寿保健"等词，会命中百岁补贴条款（如浙江"一百周岁…长寿保健补助费"），
         # 此时年龄被 99 上限排除，若再兜底成 60 就是彻头彻尾的假数据。
         allow_fallback=False,
         anchors_any=["高龄津贴", "高龄补贴", "高龄老人", "长寿保健", "高龄生活补贴",
                      "高龄老人生活补贴",
                      "敬老金", "老年年金", "生活津贴", "特别照顾津贴", "老人津贴"]),
    dict(key="freeTransitAge", label="免费乘公共交通年龄", mode="age",
         anchors_any=["公共交通", "公交车", "公共汽车", "公交", "轨道交通", "客运",
                      "巴士", "捷运", "地铁", "公车",
                      "集体运输", "大众运输", "运输工具", "公共运输"],
         anchors_all2=["免费", "优惠", "优待", "半价", "减免", "票价优惠", "豁免"]),
    dict(key="freeParkAge", label="免费进公园景区年龄", mode="age",
         anchors_any=["公园", "风景名胜区", "景区", "旅游景点", "公共文化设施",
                      "博物馆", "纪念馆", "文化馆", "体育场馆", "图书馆",
                      "风景区", "文教设施", "康乐设施", "康乐", "体育设施",
                      "文化设施", "康乐场所"],
         anchors_all2=["免费", "优惠", "优待", "免票", "半价", "豁免"]),
    dict(key="medicalPriorityAge", label="就医优待年龄", mode="age",
         anchors_any=["优先就诊", "就医优先", "绿色通道", "优先挂号", "优先就医",
                      "免费医疗", "医疗优惠"]),
    dict(key="elderlyDefAge", label="老年人定义年龄", mode="age",
         anchors_any=["老年人是指", "所称老年人", "老年人系指",
                      "长者是指", "所称长者", "长者指",
                      "老人是指", "所称老人", "老人指",
                      # 英文版（香港《安老院条例》Cap.459 等）
                      "attained the age of", "elderly person means"]),
    dict(key="careSubsidy", label="失能/养老服务补贴", mode="bool",
         anchors_any=["养老服务补贴", "护理补贴", "照护补贴", "失能补贴",
                      "长期照顾", "照顾津贴", "居家式服务", "社区式服务",
                      "机构式服务",
                      # 英文版（香港 Cap.459 规管的安老院/护养院即机构式照顾服务）
                      "residential care", "care and attention home",
                      "nursing home"]),
    dict(key="ltcInsurance", label="长期护理保险", mode="bool",
         anchors_any=["长期护理保险", "长期照护保险", "长期照顾保险"]),
    dict(key="accidentInsurance", label="老年人意外险", mode="bool",
         anchors_any=["意外伤害保险", "意外险", "意外伤害综合保险"]),
    dict(key="homeAdaptation", label="适老化改造", mode="bool",
         anchors_any=["适老化改造", "无障碍改造", "居家适老化", "家庭适老化",
                      "住家环境改善", "家居环境改善", "无障碍环境"]),
    dict(key="centenarianAllowance", label="百岁老人补贴", mode="money",
         anchors_any=["百岁", "一百周岁", "长寿保健补助", "长寿老人"]),
    dict(key="legalAid", label="法律援助", mode="bool",
         anchors_any=["法律援助", "法律支援"]),
    dict(key="filialVisit", label="赡养探望条款", mode="bool",
         anchors_any=["定期探望", "常回家看看", "与老年人分开居住", "看望老年人",
                      "关怀访视", "电话问安"]),
]


def match_sentences(sents, rule):
    out = []
    for s in sents:
        if rule.get("anchors_all") and not all(a in s for a in rule["anchors_all"]):
            continue
        if rule.get("anchors_all2") and not any(a in s for a in rule["anchors_all2"]):
            continue
        if rule.get("anchors_any") and not any(a in s for a in rule["anchors_any"]):
            continue
        out.append(s)
    return out


def extract_other(sents, rule):
    hits = match_sentences(sents, rule)
    for s in hits:
        if rule["mode"] == "age":
            # 高龄津贴排除百岁以上（那是百岁补贴）
            lim = 99 if rule["key"] == "highAgeAllowanceAge" else 200
            v = extract_age(s, max_val=lim)
            if v is not None:
                return v, s[:110]
        elif rule["mode"] == "bool":
            return True, s[:110]
        elif rule["mode"] == "money":
            m = re.search(r"(\d+)\s*元", s)
            if m:
                # 判定发放周期：先剔除"老年"再判断，避免"老年人"里的"年"误判
                ctx = s[max(0, m.start() - 14): m.end() + 4].replace("老年", "")
                if "每月" in ctx or ("月" in ctx and "年" not in ctx):
                    period = "月"
                elif "每年" in ctx or "年" in ctx:
                    period = "年"
                else:
                    period = "月"
                return f"{m.group(1)}元/{period}", s[:110]
            return True, s[:110]
    if hits and rule["mode"] == "bool":
        return True, hits[0][:110]
    return None, (hits[0][:110] if hits else "")


def extract_age_with_fallback(sents, rule, def_age):
    """
    年龄型维度抽取，带兜底：
    港澳台的福利条款常只写"长者/老人享有优惠"而不重复年龄，
    年龄隐含在"长者是指年满六十五岁"这条定义里。
    此时用该地区的老年人定义年龄作为优待起始年龄，并在证据里明确标注为推导值。
    """
    v, ev = extract_other(sents, rule)
    if v is not None or not def_age:
        return v, ev
    if rule.get("allow_fallback") is False:
        return None, ev
    # 兜底：命中了福利句但没写年龄，且句中使用了老年人称谓
    for s in match_sentences(sents, rule):
        if any(t in s for t in ELDERLY_TERMS):
            return def_age, f"{s[:80]} 〔条款未写明年龄，按当地老年定义 {def_age} 岁推导〕"
    return None, ev


# ---------------- 主流程 ----------------

def main():
    result = {}
    provinces = []
    for cat_dir in sorted(ASSETS.iterdir()):
        if not cat_dir.is_dir():
            continue
        for prov in sorted(cat_dir.iterdir()):
            if prov.is_dir():
                provinces.append((cat_dir.name, prov))

    print(f"共发现 {len(provinces)} 个省级目录\n")

    for cat, prov in provinces:
        geo = prov.name
        texts = {}
        for d in sorted(prov.iterdir()):
            if d.is_file():
                t = read_any(d)
                if t and len(t) > 200:
                    texts[d.name] = t

        all_text = "\n".join(texts.values())
        # 切分前先压缩多个空白为 1 个空格，**保留空格**——英文/港英版需依赖空格
        # 分词（"attained the age of" 不能拼成 "attainedtheageof"），中文本来就没
        # 多余空格，保留也不影响中文 anchor 匹配。
        sents = [s for s in re.split(r"[。；\n]", re.sub(r"[ \t\r]+", " ", all_text))
                 if len(s.strip()) >= 4]

        rec = {"geoName": geo, "category": cat, "docs": list(texts.keys()),
               "values": {}, "evidence": {}}
        if not texts:
            rec["values"]["_noText"] = True

        # 护理假
        cl = extract_care_leave(texts)
        rec["values"]["careLeaveDays"] = cl["days"]
        rec["values"]["careLeaveGeneralDays"] = cl["daysGeneral"]
        rec["values"]["careLeavePaid"] = cl["paid"]
        rec["values"]["careLeaveStatus"] = cl["status"]
        if cl["evidence"]:
            rec["evidence"]["careLeaveDays"] = cl["evidence"]

        # 先算"老年人定义年龄"——它同时作为各优待年龄的兜底依据
        def_rule = next(r for r in DIM_RULES if r["key"] == "elderlyDefAge")
        def_age, def_ev = extract_other(sents, def_rule)
        rec["values"]["elderlyDefAge"] = def_age
        if def_ev:
            rec["evidence"]["elderlyDefAge"] = def_ev

        # 其余维度：年龄型走兜底逻辑，其它走原逻辑
        for rule in DIM_RULES:
            if rule["key"] == "elderlyDefAge":
                continue
            if rule["mode"] == "age":
                v, ev = extract_age_with_fallback(sents, rule, def_age)
            else:
                v, ev = extract_other(sents, rule)
            rec["values"][rule["key"]] = v
            if ev:
                rec["evidence"][rule["key"]] = ev

        n = sum(1 for k, v in rec["values"].items()
                if v not in (None, False) and not k.startswith("_"))
        st = {"explicit": "有天数", "timeonly": "仅陪护时间", "none": "未规定"}[cl["status"]]
        print(f"[{n:>2}] {geo:<12} 护理假={str(cl['days'] or '-'):>3}({st}) "
              f"带薪={str(cl['paid']):<5} 高龄={str(rec['values'].get('highAgeAllowanceAge') or '-'):>3} "
              f"交通={str(rec['values'].get('freeTransitAge') or '-'):>3} "
              f"公园={str(rec['values'].get('freeParkAge') or '-'):>3} "
              f"定义={str(rec['values'].get('elderlyDefAge') or '-'):>3}")
        result[geo] = rec

    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n已输出：{OUT}")


if __name__ == "__main__":
    main()
